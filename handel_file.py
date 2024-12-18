import cloudinary
import cloudinary.uploader
from fastapi import UploadFile
from PIL import Image, ImageTk
import docx
from docx.shared import Inches,Cm,RGBColor,Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH 
import os
from datetime import datetime as dt
import datetime
import requests
from io import BytesIO
# Configure Cloudinary
cloudinary.config(
    cloud_name="dsm1ingy6",
    api_key="815961248834572",
    api_secret="KDnxt9IF0rUEXresLszPVwRqFhA"   
)
def animal_upload(File: UploadFile):
    try:
        # Read the file's content
        file_content = File.file

        # Upload the file to Cloudinary
        upload_result = cloudinary.uploader.upload(
            file_content,
            folder="Paltu", 
            resource_type="auto"
        )

        # Extract the URL from the upload result
        return upload_result["secure_url"]
    except Exception as e:
        print(f"Error uploading to Cloudinary: {e}")
        raise e
def adoptor_upload(File: UploadFile):
    try:
        # Read the file's content
        file_content = File.file

        # Upload the file to Cloudinary
        upload_result = cloudinary.uploader.upload(
            file_content,
            folder="Adopters", 
            resource_type="auto"
        )

        # Extract the URL from the upload result
        print(upload_result)
        return upload_result["secure_url"]
    except Exception as e:
        print(f"Error uploading to Cloudinary: {e}")
        raise e
def get_img_from_url(url):
    responce=requests.get(url)
    return responce.content
def Write_form(form_data: dict):
    adoptor_image=get_img_from_url(form_data["adopter_image"])
    adoptor_doc=get_img_from_url(form_data["adopter_doc"])
    animal=get_img_from_url(form_data["photos"])
    with open('images/adoptor.jpg', 'wb') as file:
                file.write(adoptor_image)
    adoptor_image='/images/adoptor.jpg'
    with open('images/doc.jpg', 'wb') as file:
                file.write(adoptor_doc)
    adoptor_doc='/images/doc.jpg'
    with open('images/animal.jpg', 'wb') as file:
                file.write(animal)
    animal='/images/animal.jpg'
    doc = docx.Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
    doc.add_paragraph(f"Serial NO - {form_data.get('serial_no')}")
    doc.add_picture('logo.jpg', width=Inches(5), height=Inches(1.5))
    img=doc.paragraphs[-1]
    img.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1=doc.add_heading("ADOPTION AND CONSENT FORM",0)
    h1.style.font.color.rgb = RGBColor(0,0,0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Councler Name: {form_data.get('counselor_name')}					Date: {datetime.date.today()} ")
    dic=doc.add_heading("DESCRIPTION OF ANIMAL",1)
    dic.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dic.style.font.color.rgb = RGBColor(0,0,0)
    animal=doc.add_table(rows=1,cols=2) 
    animal.cell(0,0).add_paragraph(f"Colour: {form_data.get('color')}        Tag:{form_data.get('tag')}\nGender: {form_data.get('gender')}    Age: {form_data.get('age')}\nBreed: {form_data.get('breed')}\nPhysically Fitness Status: {form_data.get('fitness_status')}\nVaccination Status: {form_data.get('vaccination_status')}\nSterilisation Status: {form_data.get('sterilisation_status')}")
    animal.cell(0,1).add_paragraph().add_run().add_picture(animal,width=Inches(2),height=Inches(1.75))
    details=doc.add_heading("DETAILS OF CARETAKER                   DETAILS OF ADOPTER",1)
    details.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    details.style.font.color.rgb = RGBColor(0,0,0)
    caretaker=doc.add_table(rows=2,cols=2)
    caretaker.cell(0,0).add_paragraph(f"Name of Caretaker: {form_data.get('caretaker_name')}\nContact No.: {form_data.get('caretaker_contact')}\nWhatsapp No.: {form_data.get('caretaker_whatsapp')}\nEmail I'd: {form_data.get('caretaker_email')}\nLocal Residence: {form_data.get('caretaker_local_residence')}\nPermanent Residence: {form_data.get('caretaker_permanent_residence')}\nInstagram/Facebook ID: {form_data.get('caretaker_social_id')}\nSigneture:_______________")
    formated_text=f"Name of Adopter: {form_data.get('adopter_name')}\nContact No.: {form_data.get('adopter_contact')}\nWhatsapp No.: {form_data.get('adopter_whatsapp')}\nEmail Id: {form_data.get('adopter_email')}\nLocal Residence: {form_data.get('adopter_local_residence')}\nPermanent Residence: {form_data.get('adopter_permanent_residence')}\nGovernment ID number: {form_data.get('adopter_gov_id')}\nInstagram/Facebook ID: {form_data.get('adopter_social_id')}\nSigneture:_______________" 
    caretaker.cell(0,1).add_paragraph(formated_text)
    caretaker.cell(1,0).add_paragraph().add_run().add_picture(adoptor_doc,width=Inches(2),height=Inches(2))
    img=caretaker.cell(1,1).add_paragraph()
    img.add_run().add_picture(adoptor_image,width=Inches(2),height=Inches(2))
    img.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
#    doc.add_heading("General information from Adopter",1)
#    para=doc.add_paragraph(f"What are your plans for your pet if you shift to another town/place?\n {form_data.get('plans')}")
#    para.add_run(f"\nHave you had a pet before or have one right now?\n{form_data.get('pet_experience')}")
#    para.add_run(f"\nAmount of time your pet may have to be alone in a day?\n{form_data.get('time_alone')}")
#    para.add_run(f"\nWho will take care of your pet if you go out of town temporarily?\n{form_data.get('caretaker_if_away')}")
    con=doc.add_heading("Consent By Adopter",1)
    con.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    con.style.font.color.rgb = RGBColor(0,0,0)
    doc.add_paragraph('With the unanimous consent of all my family members, I am willingly adopting the pet, assuming full responsibility and acknowledging that I will always regard my pet as an integral part of our family, ensuring that it is never treated merely as an object.',style='List Bullet')
    doc.add_paragraph("I commit to maintaining a clean and well-ventilated environment for the adopted pet, providing proper nourishment, and ensuring regular exercise. I will refrain from keeping the animal tethered or chained with an unreasonably short or heavy restraint for an extended duration.",style='List Bullet')
    doc.add_paragraph("I acknowledge that certain diseases may not be detectable during the initial checkup of the adopted pet, and the Animals With Humanity Team will not be held responsible for such conditions. Therefore, it is recommended to conduct a comprehensive post-adoption health examination.",style='List Bullet')
    doc.add_paragraph("If the pet I adopt becomes unwell, I will promptly seek advice from a veterinarian and notify the Animals With Humanity Team. Additionally, I will take responsibility for adhering to the deworming, vaccination, and sterilization schedule for the well-being of the pet.",style='List Bullet')
    doc.add_paragraph("Should the responsibility of pet parenting need to be transferred, I will provide advance notice to both the caretaker and Team Animals With Humanity. I understand that the adoption process will need to be repeated for the new caretakers, and a fine of Rs. 5,000/- will be duly acknowledged and paid.",style='List Bullet')
    doc.add_paragraph("I acknowledge that the Animals With Humanity Team possesses the authority to conduct unannounced inspections of the conditions in which I'm taking care for the adopted pet. In the event of any violations, the Animals With Humanity Team is empowered to take legal action as per applicable law.",style='List Bullet')
    doc.add_paragraph(f"I {form_data.get('adopter_name')} acknowledge that abandoning or subjecting my pet to mistreatment may lead to legal consequences under the Prevention of Cruelty to Animals Act of 1960. In case my pet is found in any such situation, Team Animals With Humanity will take a fine upto Rs. 10,000/- depending on the situation and proceed with legal action.",style='List Bullet')
    doc.add_paragraph("I enter into this contract of my own free will and understand that this is a binding contract enforceable by civil law.",style='List Bullet')
    doc.save(self.file_path)